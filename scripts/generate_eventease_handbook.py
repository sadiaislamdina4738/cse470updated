"""
Generates EventEase_Handbook.docx in the project root (cse470/).
Run: python scripts/generate_eventease_handbook.py
"""
from pathlib import Path

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_LINE_SPACING


def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    return p


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    return h


def main():
    root = Path(__file__).resolve().parent.parent
    out_path = root / "EventEase_Handbook.docx"

    doc = Document()
    core = doc.core_properties
    core.title = "EventEase Handbook"
    core.author = "EventEase"

    doc.add_heading("EventEase — Project handbook", 0)

    add_heading(doc, "1. Admin test account (local demo)", 1)
    add_para(
        doc,
        "Use this account to sign in on the website (or register if it does not exist yet). "
        "Change the password after coursework if you share this document.",
    )
    doc.add_paragraph()
    t = doc.add_table(rows=4, cols=2)
    t.style = "Table Grid"
    rows = [
        ("Field", "Value"),
        ("Email", "sadia.islam9@g.bracu.ac.bd"),
        ("Username", "dina"),
        ("Password", "dina@1234"),
    ]
    for i, (a, b) in enumerate(rows):
        t.rows[i].cells[0].text = a
        t.rows[i].cells[1].text = b
    add_para(
        doc,
        "Admin dashboard: after login, if this user has role admin, open Admin in the navbar or go to http://localhost:3000/admin . "
        "Admin is granted on first registration when email matches INITIAL_ADMIN_EMAIL in backend/.env .",
    )

    add_heading(doc, "2. Backend environment (backend/.env)", 1)
    add_para(
        doc,
        "Create backend/.env from backend/.env.example . Never commit real secrets to Git. "
        "Below is the meaning of each variable (values are the ones your project uses locally; rotate passwords if exposed).",
    )
    env_lines = [
        ("MONGODB_URI", "MongoDB connection string (Atlas standard URI or mongodb+srv). Keep secret."),
        ("JWT_SECRET", "Secret string used to sign JWT tokens."),
        ("PORT", "API port (default 5000)."),
        ("NODE_ENV", "development or production."),
        ("FRONTEND_URL", "Allowed CORS origin, e.g. http://localhost:3000 ."),
        (
            "INITIAL_ADMIN_EMAIL",
            "Optional. First successful registration with this exact email gets role admin.",
        ),
    ]
    et = doc.add_table(rows=1 + len(env_lines), cols=2)
    et.style = "Table Grid"
    et.rows[0].cells[0].text = "Variable"
    et.rows[0].cells[1].text = "Description"
    for i, (k, desc) in enumerate(env_lines, start=1):
        et.rows[i].cells[0].text = k
        et.rows[i].cells[1].text = desc

    add_para(
        doc,
        "Example .env layout (replace placeholders with your own secrets):",
        bold=True,
    )
    add_para(
        doc,
        "MONGODB_URI=mongodb://USER:PASSWORD@host1:27017,host2:27017,host3:27017/eventease?ssl=true&authSource=admin&replicaSet=YOUR_REPLICA_SET&retryWrites=true&w=majority\n"
        "JWT_SECRET=your_long_random_secret\n"
        "PORT=5000\n"
        "NODE_ENV=development\n"
        "FRONTEND_URL=http://localhost:3000\n"
        "INITIAL_ADMIN_EMAIL=sadia.islam9@g.bracu.ac.bd",
    )

    env_path = root / "backend" / ".env"
    if env_path.is_file():
        add_heading(doc, "2b. Your backend/.env keys (secrets redacted)", 2)
        add_para(
            doc,
            "The lines below mirror backend/.env on disk. Database user/password and JWT value are not "
            "printed in full—open backend/.env for the real values. Do not share that file publicly.",
        )
        for raw in env_path.read_text(encoding="utf-8", errors="replace").splitlines():
            line = raw.strip()
            if not line or line.startswith("#"):
                continue
            if line.upper().startswith("MONGODB_URI="):
                add_para(doc, "MONGODB_URI=<redacted — full URI in backend/.env>")
            elif line.upper().startswith("JWT_SECRET="):
                add_para(doc, "JWT_SECRET=<redacted — see backend/.env>")
            else:
                add_para(doc, line)

    add_heading(doc, "3. Implemented features (summary)", 1)
    features = [
        "Authentication: register, login, JWT, /api/auth/me, profile; deactivated users blocked.",
        "Events: create, list with filters (category, location, keyword q, date range from/to, sort), map and calendar views on list page.",
        "Event detail: view, RSVP / pending / waitlist, cancel RSVP, edit/delete for creator.",
        "Organizer: manage attendees (approve, reject, remove, promote waitlist), toggle requiresApproval.",
        "Real-time: Socket.io event-updated for RSVP/attendee changes; chat rooms separate from live updates.",
        "Chat: REST + Socket.io; access for creator, attendees, pending, waitlist; route order fixed for GET /api/events/chat .",
        "QR / share: canonical URL from GET /api/events/:id/share-link ; Share tab with QR and copy link.",
        "Feedback: after event date and only if attended; rating distribution on display.",
        "Profile: account tab + My calendar (hosted + attended events) with day details and links.",
        "Admin (role admin): /admin — users (search, deactivate/reactivate), events (deactivate), analytics.",
        "Security: CORS restricted to FRONTEND_URL; rate limit on login/register.",
    ]
    for f in features:
        doc.add_paragraph(f, style="List Bullet")

    add_heading(doc, "4. How to access in the website", 1)
    add_para(doc, "Run backend and frontend first (section 5).", bold=True)
    urls = [
        ("Home", "http://localhost:3000/"),
        ("Register / Login", "http://localhost:3000/register , http://localhost:3000/login"),
        ("Browse events (discovery)", "http://localhost:3000/events"),
        ("Create event", "http://localhost:3000/create-event (logged in)"),
        ("Profile + calendar", "http://localhost:3000/profile — tabs Account / My calendar"),
        ("Admin dashboard", "http://localhost:3000/admin (admin user only)"),
        ("API base (reference)", "http://localhost:5000/api/..."),
    ]
    for name, url in urls:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(f"{name}: ").bold = True
        p.add_run(url)

    add_para(
        doc,
        "Typical flow: Register → Events → open an event → RSVP → use tabs Details, Manage (if creator), Chat, Feedback, Share.",
    )

    add_heading(doc, "5. How to run from a folder (Windows)", 1)
    add_para(doc, "Assume the project folder is:", bold=True)
    add_para(doc, r"F:\project cse 470\cse470")
    add_para(doc, "Terminal 1 — backend:", bold=True)
    add_para(
        doc,
        r'cd "F:\project cse 470\cse470\backend"'
        "\n"
        "npm install\n"
        "npm run dev\n"
        "# Wait for: Server running on port 5000 and Connected to MongoDB",
    )
    add_para(doc, "Terminal 2 — frontend:", bold=True)
    add_para(
        doc,
        r'cd "F:\project cse 470\cse470\frontend"'
        "\n"
        "npm install\n"
        "npm start\n"
        "# Browser opens http://localhost:3000",
    )
    add_para(
        doc,
        "Optional: copy backend/.env.example to backend/.env and fill MONGODB_URI and JWT_SECRET before npm run dev .",
    )

    add_heading(doc, "6. Production build (optional)", 1)
    add_para(doc, "Frontend: cd frontend && npm run build — output in frontend/build .")
    add_para(doc, "Serve build with any static host; set REACT_APP_API_URL to your deployed API.")

    doc.save(out_path)
    print(f"Wrote: {out_path}")


if __name__ == "__main__":
    main()
